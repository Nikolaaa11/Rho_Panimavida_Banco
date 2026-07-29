# -*- coding: utf-8 -*-
"""Brochure Panimavida v8 en Word editable + correo para Javier."""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import os as _os
_BASE = _os.path.dirname(_os.path.abspath(__file__))
_DATOS = _os.path.join(_BASE, 'datos')


INK = RGBColor(0x1d, 0x1d, 0x1f)
GREY = RGBColor(0x6e, 0x6e, 0x73)
GREYL = RGBColor(0x86, 0x86, 0x8b)
GREEN = RGBColor(0x21, 0x83, 0x58)
RED = RGBColor(0xe5, 0x48, 0x4d)
AMBER = RGBColor(0xb2, 0x5e, 0x09)
FONT = 'Segoe UI'

OUTDIR = r'C:\Users\nicol\OneDrive\Documentos\0.2.Rho\Banco_Panimavida\entrega_banco_v8'
os.makedirs(OUTDIR, exist_ok=True)
CHART = _os.path.join(_DATOS, 'bro_chart.png')


def setfont(doc):
    st = doc.styles['Normal']
    st.font.name = FONT
    st.font.size = Pt(10.5)
    st.font.color.rgb = INK
    st.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    pf = st.paragraph_format
    pf.space_after = Pt(7)
    pf.line_spacing = 1.22


def p(doc, text='', size=10.5, bold=False, color=INK, italic=False,
      before=0, after=7, align=None, indent=None):
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(before)
    par.paragraph_format.space_after = Pt(after)
    if align:
        par.alignment = align
    if indent:
        par.paragraph_format.left_indent = Cm(indent)
    if text:
        r = par.add_run(text)
        r.font.name = FONT
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.italic = italic
        r.font.color.rgb = color
    return par


def rich(doc, parts, size=10.5, before=0, after=7, indent=None):
    """parts = [(texto, bold, color), ...]"""
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(before)
    par.paragraph_format.space_after = Pt(after)
    if indent:
        par.paragraph_format.left_indent = Cm(indent)
    for t, b, c in parts:
        r = par.add_run(t)
        r.font.name = FONT
        r.font.size = Pt(size)
        r.font.bold = b
        r.font.color.rgb = c
    return par


def eyebrow(doc, text):
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(3)
    r = par.add_run(' '.join(text.upper()))
    r.font.name = FONT
    r.font.size = Pt(7.5)
    r.font.bold = True
    r.font.color.rgb = GREEN


def h1(doc, text):
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(2)
    par.paragraph_format.space_after = Pt(9)
    r = par.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(21)
    r.font.color.rgb = INK


def h2(doc, text):
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(13)
    par.paragraph_format.space_after = Pt(5)
    r = par.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(12.5)
    r.font.bold = True
    r.font.color.rgb = INK


def bullet(doc, text, color=INK, size=10):
    rich(doc, [('·  ', False, GREEN), (text, False, color)], size=size, after=5, indent=0.4)


def shade(cell, hexcolor):
    el = OxmlElement('w:shd')
    el.set(qn('w:val'), 'clear')
    el.set(qn('w:fill'), hexcolor)
    cell._tc.get_or_add_tcPr().append(el)


def celltxt(cell, text, size=9.5, bold=False, color=INK, align=None, italic=False):
    cell.text = ''
    par = cell.paragraphs[0]
    par.paragraph_format.space_before = Pt(3)
    par.paragraph_format.space_after = Pt(3)
    if align:
        par.alignment = align
    r = par.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = color


def mktable(doc, rows, cols, widths=None, style='Table Grid'):
    t = doc.add_table(rows=rows, cols=cols)
    t.style = style
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    if widths:
        for r_ in t.rows:
            for i, w in enumerate(widths):
                r_.cells[i].width = Cm(w)
    return t


R = WD_ALIGN_PARAGRAPH.RIGHT
C_ = WD_ALIGN_PARAGRAPH.CENTER

# ══════════════════════════════════════════════════════ BROCHURE
doc = Document()
setfont(doc)
for s in doc.sections:
    s.top_margin = Cm(2.0)
    s.bottom_margin = Cm(1.8)
    s.left_margin = Cm(2.1)
    s.right_margin = Cm(2.1)

# ---------- Pagina 1
eyebrow(doc, 'Panimávida 13,2 kV  ·  análisis para el comité de crédito  ·  julio 2026')
h1(doc, 'No está cayendo lo que se cree que está cayendo.')
p(doc, 'La objeción planteada al proyecto es que la línea HVDC Kimal–Lo Aguirre hará bajar el precio '
       'de venta nocturno del nodo. Los datos del propio nodo, medidos en ventanas de doce meses '
       'completos, muestran otra cosa.', 10.5, color=GREY, after=11)

p(doc, 'Ventanas de 12 meses completos · 1.586 días reales del Coordinador Eléctrico Nacional',
  9, bold=True, after=4)
t = mktable(doc, 5, 4, [5.6, 3.4, 3.6, 3.0])
for i, lb in enumerate(['Ventana', 'Costo de carga', 'Precio de venta', 'Spread']):
    celltxt(t.rows[0].cells[i], lb, 8.5, True, GREY, None if i == 0 else R)
    shade(t.rows[0].cells[i], 'F7F9F8')
datos = [('2022 – 2023', '37,7', '200,0', '162,3', False),
         ('2023 – 2024', '12,3', '108,2', '96,0', False),
         ('2024 – 2025', '10,1', '98,2', '88,1', False),
         ('2025 – 2026', '22,1', '99,6', '77,5', True)]
for i, (lb, cg, vt, sp, last) in enumerate(datos, 1):
    celltxt(t.rows[i].cells[0], lb, 9.5, last, INK)
    celltxt(t.rows[i].cells[1], cg, 9.5, last, RED if last else INK, R)
    celltxt(t.rows[i].cells[2], vt, 9.5, True, GREEN, R)
    celltxt(t.rows[i].cells[3], sp, 9.5, last, INK, R)
    if last:
        for k in range(4):
            shade(t.rows[i].cells[k], 'F2F9F5')
p(doc, 'USD/MWh · ventana de 4 horas · carga 09–17 h, venta 19–07 h', 8, color=GREYL,
  italic=True, before=3, after=12)

p(doc, 'El precio de venta nocturno está plano desde hace tres años.', 12.5, True, GREEN, after=4)
rich(doc, [('108,2 → 98,2 → 99,6 USD/MWh. ', True, INK),
           ('Lo que se movió, y explica prácticamente toda la compresión reciente del spread, es el '
            'costo de cargar de día: subió de 10,1 a 22,1 USD/MWh.', False, INK)], after=12)

h2(doc, 'El giro del argumento')
bullet(doc, 'Kimal–Lo Aguirre transporta excedente solar. El sol es un recurso diurno; la batería '
            'vende de noche. La línea actúa sobre el costo marginal del día —es decir, sobre el lado '
            'donde está el riesgo real— y lo empuja a la baja.')
bullet(doc, 'Sobre la evidencia disponible, la línea que se teme es viento de cola para este '
            'proyecto, no viento de frente. Y no llega al Maule: termina en Lo Aguirre, Región '
            'Metropolitana. Su entrada en operación se estima en 2029–2030.')

h2(doc, 'El piso del precio de venta')
p(doc, 'El caso no se sostiene en una cifra única, sino en la distribución completa del registro: en '
       '1.586 días, el precio de venta de las cuatro horas de punta estuvo en o sobre 50 USD/MWh el '
       '99,2% de los días.', 10, after=6)
t = mktable(doc, 2, 4, [3.9, 3.9, 3.9, 3.9])
for i, (pc, lb) in enumerate([('99,7%', 'en o sobre 40 USD/MWh'), ('99,2%', 'en o sobre 50 USD/MWh'),
                              ('98,1%', 'en o sobre 55 USD/MWh'), ('94,5%', 'en o sobre 60 USD/MWh')]):
    acc = i == 1
    celltxt(t.rows[0].cells[i], pc, 15, True, GREEN if acc else INK, C_)
    celltxt(t.rows[1].cells[i], lb, 8, False, GREY, C_)
    if acc:
        shade(t.rows[0].cells[i], 'F2F9F5')
        shade(t.rows[1].cells[i], 'F2F9F5')
p(doc, 'El mínimo absoluto del registro es 0,0 USD/MWh (30-oct-2022), un día atípico. En los últimos '
       '12 meses el peor día fue 48,8 y el percentil 10 fue 58,5 USD/MWh.', 8, color=GREYL,
  italic=True, before=4)

doc.add_page_break()

# ---------- Pagina 2
eyebrow(doc, 'la ventana de arbitraje  ·  2027 – 2035')
h1(doc, 'Hasta cuándo conviene el marginal.')
p(doc, 'Una sola variable manda: la velocidad con que el almacenamiento del sistema llega a comprimir '
       'este nodo. Se modela por los dos lados —la venta baja y la carga sube— y se compara contra el '
       'spread mínimo necesario para pagar la deuda.', 10.5, color=GREY, after=11)

t = mktable(doc, 3, 3, [5.2, 5.2, 5.2])
kpi = [('815', 'mil US$ / año', 'Margen bruto de los últimos 12 meses reales · 340 ciclos', True),
       ('43,9', 'USD/MWh', 'Spread necesario para cubrir el servicio de la deuda', False),
       ('2034', 'escenario base', 'Año en que el spread cruza ese umbral', False)]
for i, (num, un, lb, acc) in enumerate(kpi):
    celltxt(t.rows[0].cells[i], num, 22, True, GREEN if acc else INK)
    celltxt(t.rows[1].cells[i], un, 8, False, GREYL)
    celltxt(t.rows[2].cells[i], lb, 8.5, False, GREY)
    if acc:
        for k in range(3):
            shade(t.rows[k].cells[i], 'F2F9F5')

h2(doc, 'Tres escenarios de compresión')
t = mktable(doc, 4, 4, [4.6, 3.5, 3.6, 4.0])
for i, lb in enumerate(['Escenario', 'Spread 2035', 'Cruza el break-even', 'Inicio de la compresión']):
    celltxt(t.rows[0].cells[i], lb, 8.5, True, GREY)
    shade(t.rows[0].cells[i], 'F7F9F8')
esc = [('A · Ventana amplia', '62,5', 'no cruza al 2035', 'desde 2031', GREEN),
       ('B · Base', '32,0', '2034', 'desde 2029', AMBER),
       ('C · Estrés tipo NEM', '10,0', '2029', 'desde 2027', RED)]
for i, (nm, sp, cr, ini, col) in enumerate(esc, 1):
    celltxt(t.rows[i].cells[0], nm, 9.5, True, INK)
    celltxt(t.rows[i].cells[1], sp, 9.5, False, INK)
    celltxt(t.rows[i].cells[2], cr, 9.5, True, col)
    celltxt(t.rows[i].cells[3], ini, 9, False, GREY)
p(doc, '', after=4)
for nm, just in [('A', 'El almacenamiento sigue concentrado en el norte y la CNE proyecta desacople del sur hacia 2032.'),
                 ('B', 'Compresión gradual desde la entrada de Kimal–Lo Aguirre (2029–2030) y llegada parcial de BESS al centro-sur.'),
                 ('C', 'Replica el precedente australiano: el spread cayó 85% en un año al pasar la flota BESS de 4.360 a 9.000 MW.')]:
    rich(doc, [(f'{nm} · ', True, GREEN), (just, False, GREY)], size=8.8, after=3, indent=0.3)

h2(doc, 'Cómo se lee la holgura')
for t_ in ['El spread observado en los últimos 12 meses completos es 77,5 USD/MWh.',
           'A 340 ciclos, cubrir un servicio de deuda de US$400 mil requiere 43,9 USD/MWh.',
           'El spread tendría que caer un 43% antes de comprometer el servicio de la deuda.',
           'En esa misma ventana, el 88% de los días superó los 40 USD/MWh de spread.']:
    bullet(doc, t_, size=9.5)
p(doc, 'Los supuestos de deuda son editables en la calculadora; ninguna cifra de DSCR es válida hasta '
       'cargar el term sheet real.', 8.2, color=GREYL, italic=True, before=3, after=10)

if os.path.exists(CHART):
    p(doc, 'Spread proyectado por escenario · la línea roja punteada es el break-even de 43,9 USD/MWh',
      8.4, True, GREY, after=4)
    doc.add_picture(CHART, width=Cm(16.8))

doc.add_page_break()

# ---------- Pagina 3
eyebrow(doc, 'qué contrato, y cuándo  ·  monitoreo')
h1(doc, 'El contrato correcto, en su momento.')
p(doc, 'Decir "no PPA ahora, reevaluar hacia 2030" suena contradictorio. No lo es: son dos contratos '
       'distintos que hoy se llaman igual.', 10.5, color=GREY, after=10)

t = mktable(doc, 6, 3, [3.9, 6.0, 6.0])
celltxt(t.rows[0].cells[0], '', 9)
celltxt(t.rows[0].cells[1], 'PPA de COMPRA a 28 USD/MWh', 10, True, RED)
celltxt(t.rows[0].cells[2], 'Instrumento de VENTA · 2030–2032', 10, True, GREEN)
shade(t.rows[0].cells[1], 'FDF2F2')
shade(t.rows[0].cells[2], 'F2F9F5')
filas = [('Qué fija', 'El precio al que el proyecto compra su energía de carga',
          'El precio al que el proyecto vende en la noche'),
         ('Riesgo que cubre', 'Que suba el precio de compra', 'Que caiga el precio de venta'),
         ('¿Es el riesgo real?', 'No. El proyecto es posición vendedora',
          'Sí. Es exactamente el riesgo del activo'),
         ('Efecto', 'Destruye cerca de US$83 mil/año de margen', 'Estabiliza el margen y la caja'),
         ('Veredicto', 'Rechazar de forma permanente, no "por ahora"',
          'Negociar desde el año 1; firmar entre 2030 y 2032')]
for i, (k, a, b) in enumerate(filas, 1):
    last = k == 'Veredicto'
    celltxt(t.rows[i].cells[0], k, 9, True, GREEN if last else GREY)
    celltxt(t.rows[i].cells[1], a, 9, last, RED if last else INK)
    celltxt(t.rows[i].cells[2], b, 9, last, GREEN if last else INK)
    if last:
        shade(t.rows[i].cells[1], 'FDF2F2')
        shade(t.rows[i].cells[2], 'F2F9F5')

h2(doc, 'Cuánto se gana con PPA y cuánto sin PPA')
p(doc, 'Resultado real del nodo en los últimos doce meses completos bajo las dos alternativas. Para '
       'que la comparación sea justa, al PPA se le concede su mejor política de operación: con un '
       'costo de carga fijo conviene ciclar todos los días, incluso aquellos en que el spread spot es '
       'delgado. Aun así pierde.', 9.8, after=6)
t = mktable(doc, 6, 4, [5.2, 3.6, 3.9, 3.1])
for i, lb in enumerate(['', 'Sin PPA · a spot', 'Con PPA a 28', 'Diferencia']):
    celltxt(t.rows[0].cells[i], lb, 8.5, True, GREY, None if i == 0 else R)
    shade(t.rows[0].cells[i], 'F7F9F8')
cmp_rows = [('Ciclos completos en el año', '340', '365', '+25', GREY),
            ('Ingreso por venta nocturna', '997.905', '1.086.101', '+88.196', GREY),
            ('Costo de la energía de carga', '183.167', '367.920', '+184.753', RED),
            ('Costo de carga implícito', '15,0 USD/MWh', '28,0 USD/MWh', '+13,0', RED)]
for i, (k, a, b, dd, dcol) in enumerate(cmp_rows, 1):
    celltxt(t.rows[i].cells[0], k, 9.2, False, INK)
    celltxt(t.rows[i].cells[1], a, 9.2, False, INK, R)
    celltxt(t.rows[i].cells[2], b, 9.2, False, INK, R)
    celltxt(t.rows[i].cells[3], dd, 9.2, False, dcol, R)
celltxt(t.rows[5].cells[0], 'Margen bruto del año', 9.6, True, INK)
celltxt(t.rows[5].cells[1], '814.738', 9.6, True, GREEN, R)
celltxt(t.rows[5].cells[2], '718.181', 9.6, True, INK, R)
celltxt(t.rows[5].cells[3], '−96.557', 9.6, True, RED, R)
for k in range(4):
    shade(t.rows[5].cells[k], 'F2F9F5')
p(doc, 'Cifras en US$ salvo donde se indique. Ventana de 12 meses completos, 23-may-2025 a '
       '22-may-2026, sobre precios reales del nodo.', 8, color=GREYL, italic=True, before=3, after=5)
rich(doc, [('Con el PPA el proyecto cicla 25 veces más y captura US$88 mil más de ingreso, pero paga '
            'US$185 mil más por la energía de carga. El saldo son US$97 mil menos de margen al año, '
            'un 11,9%. Sobre diez años de deuda, US$966 mil.', True, INK)], size=9.6, after=9)
p(doc, 'Y la conclusión no depende del precio ofertado: evaluado entre 22 y 34 USD/MWh, y dándole a '
       'cada caso su política óptima, el PPA de compra no gana en ningún punto del rango. A 22 USD/MWh '
       'el spot todavía aventaja en US$18 mil al año; a 34, en US$175 mil.', 9.8, after=10)

h2(doc, 'Por qué la ventana existe: la brecha es locacional, no temporal')
p(doc, 'El almacenamiento en Chile es real y avanza rápido: 2.283 MW en operación y 6.358 MW en '
       'construcción. Pero está concentrado en el norte —Antofagasta 44%, Atacama 23%, Tarapacá '
       '13,6%— y el Maule no figura en el desglose oficial. La Comisión Nacional de Energía proyecta '
       'que el desacople de costos marginales de la zona sur persiste hacia 2032 incluso con las '
       'obras de 500 kV previstas.', 10)

h2(doc, 'La decisión se monitorea, no se apuesta')
p(doc, 'Seis gatillos con métrica, fuente, umbral y acción predefinida, revisados trimestralmente: '
       'compresión del spread, erosión del piso nocturno, alza del costo de carga, MW de '
       'almacenamiento en operación en el centro-sur, entrada de obras de transmisión y cobertura de '
       'deuda. El precedente australiano —spread 85% menor en un año— es la razón por la que la '
       'revisión es trimestral y no anual.', 10)

p(doc, 'FUENTES', 8, True, GREY, before=12, after=3)
for s_ in ['Coordinador Eléctrico Nacional, Real Definitivo · nodo BA S/E Panimávida 13,2 kV (BP1) · '
           '38.064 horas, 2022–2026',
           'Comisión Nacional de Energía · Informe Final del Plan de Expansión de la Transmisión 2026',
           'CNE · Norma Técnica de Conexión y Operación de PMGD, 19-feb-2026 · ACERA · Energy-Storage.News']:
    p(doc, s_, 7.8, color=GREYL, after=2)

out_doc = os.path.join(OUTDIR, 'Panimavida_Brochure_v8_EDITABLE.docx')
doc.save(out_doc)
print('OK ->', out_doc, os.path.getsize(out_doc), 'bytes')
